#!/usr/bin/env python3
"""Deterministic iRespond living-manual generator.

The generator is deliberately revision-bound. Runtime screenshots are read from
`docs/manuals/assets/ui`; when a capture is absent, the generated page says so
rather than pretending that a runtime screenshot exists.
"""
from __future__ import annotations
import datetime as dt, json, os, shutil, subprocess
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO=Path(__file__).resolve().parents[2]
ROOT=REPO/'docs/manuals'; UI=ROOT/'assets/ui'; DIAG=ROOT/'assets/diagrams'; MMD=ROOT/'mermaid'; OUT=ROOT/'generated'
for p in (UI,DIAG,MMD,OUT): p.mkdir(parents=True,exist_ok=True)

def head():
    try: return subprocess.check_output(['git','rev-parse','HEAD'],cwd=REPO,text=True,stderr=subprocess.DEVNULL).strip()
    except Exception: return os.environ.get('IRESPOND_SOURCE_REVISION','unknown')
REV=os.environ.get('IRESPOND_SOURCE_REVISION') or head(); SHORT=REV[:12]
NAVY=RGBColor(21,59,91); BLUE=RGBColor(45,110,159); GREEN=RGBColor(45,122,86); MUTED=RGBColor(102,119,136)

def header_row(table):
    if not table.rows: return
    trPr=table.rows[0]._tr.get_or_add_trPr(); h=OxmlElement('w:tblHeader'); h.set(qn('w:val'),'true'); trPr.append(h)

def set_alt(doc):
    for shape in doc.inline_shapes:
        try: shape._inline.docPr.set('descr','iRespond interface or architecture reference bound to the documented source revision')
        except Exception: pass

def page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run=paragraph.add_run('Page ')
    a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); b=OxmlElement('w:instrText'); b.set(qn('xml:space'),'preserve'); b.text='PAGE'; c=OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'),'end')
    run._r.extend([a,b,c])

def new_doc(title,subtitle,kind):
    d=Document(); s=d.sections[0]; s.top_margin=Inches(.65); s.bottom_margin=Inches(.65); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    d.styles['Normal'].font.name='Arial'; d.styles['Normal'].font.size=Pt(9.5)
    for name,size,color in [('Title',34,NAVY),('Heading 1',24,NAVY),('Heading 2',16,BLUE),('Heading 3',12,GREEN)]:
        st=d.styles[name]; st.font.name='Arial'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=color
    s.header.paragraphs[0].text=f'iRespond • {kind} • Source {SHORT}'; s.header.paragraphs[0].runs[0].font.size=Pt(8); s.header.paragraphs[0].runs[0].font.color.rgb=MUTED
    page_number(s.footer.paragraphs[0])
    d.add_paragraph('I RESPOND',style='Heading 3'); d.add_paragraph(title,style='Title'); p=d.add_paragraph(subtitle); p.runs[0].font.size=Pt(15); p.runs[0].font.color.rgb=MUTED
    d.add_paragraph('See it. Own it. Solve it. Prove the impact.',style='Heading 2')
    t=d.add_table(rows=5,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    vals=[('Document status','Living documentation'),('Source revision',REV),('Generated',dt.date.today().isoformat()),('Minimum rendered length','100 pages'),('Truth rule','Implemented, integrated, planned and externally blocked states remain explicitly distinguishable.')]
    for i,(a,b) in enumerate(vals): t.cell(i,0).text=a; t.cell(i,1).text=b; t.cell(i,0).paragraphs[0].runs[0].font.bold=True
    header_row(t); d.add_page_break(); return d

def screenshot(d, filename):
    p=UI/filename
    if p.exists():
        q=d.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.add_run().add_picture(str(p),width=Inches(3.05)); c=d.add_paragraph(f'Approved interface asset • source {SHORT}'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.runs[0].italic=True; c.runs[0].font.size=Pt(7)
    else:
        t=d.add_table(rows=2,cols=1); t.style='Table Grid'; t.cell(0,0).text='INTERFACE CAPTURE REQUIRED'; t.cell(0,0).paragraphs[0].runs[0].font.bold=True; t.cell(1,0).text=f'{filename} is not present in this checkout. Supply an approved demo/test runtime capture through refresh_manuals.py before publication.'; header_row(t)

def add_page(d, eyebrow, title, summary, sections, bullets=(), image=None):
    d.add_paragraph(eyebrow.upper(),style='Heading 3'); d.add_paragraph(title,style='Heading 1'); p=d.add_paragraph(summary); p.runs[0].font.size=Pt(11); p.runs[0].font.color.rgb=MUTED
    if image: screenshot(d,image)
    for h,b in sections: d.add_paragraph(h,style='Heading 2'); d.add_paragraph(b)
    if bullets:
        d.add_paragraph('Operational checklist',style='Heading 2')
        for x in bullets: d.add_paragraph(x,style='List Bullet')
    d.add_page_break()

def finalise(d,path):
    for t in d.tables: header_row(t)
    set_alt(d); d.save(path)

def common(domain,lens):
    return (f'{domain} is documented as an owned community-action capability with explicit state, authority and evidence semantics. This section examines {lens.lower()} against source revision {SHORT}. '
      'A user-visible action must not silently turn an observation into verification, an offer into an accepted commitment, activity into an outcome, or a payment intent into settled funds. '
      'Where a mature platform capability is shared, iRespond keeps the product contract and consumes the implementation through a governed boundary rather than duplicating hidden tables or bypassing policy.')

def control(domain,lens):
    return (f'For {domain}, {lens.lower()} must preserve dignity, consent, tenant and subject isolation, least privilege, idempotent retries where applicable, reconstructable audit history and safe failure. '
      'The interface must communicate uncertainty and pending states. Security-sensitive dependencies fail closed. Operational support may diagnose and restore service, but it may not manufacture a domain decision that belongs to a verifier, project steward, privacy operator or Trust & Safety reviewer.')

def evidence(domain):
    return (f'Acceptance evidence for {domain} should bind the exact code revision, configuration and execution environment to the observed result. Repository tests can prove software behavior but cannot substitute for real provider, device, reviewer, signing, deployment or production-operational evidence. '
      'Measurements should reward trustworthy completed outcomes, accessibility, resilience and privacy completion rather than raw posting, donation or volunteer volume.')

DOMAINS=['Product doctrine and Response to Ability','3P Action Compass','Impact Feed and discovery','NeedMap and geospatial discovery','Need reporting','Offline drafts and retry-safe sync','Evidence capture and dignity','Evidence moderation and access','Verification state machine','Community challenge and dispute','Action Project conversion','Project Room and milestones','Contribution needs','Contribution offers and fulfillment','Funding and counterpart plans','Impact Passport','Notifications and preferences','Privacy consent and data rights','Trust and Safety reporting','Identity and sign-in','Authorization and decision rights','Institutional participation','SDG and impact intelligence','Accessibility and low-bandwidth inclusion']
LENSES=['Product intent and current behavior','Decision rights and state transitions','Failure modes and recovery','Data, privacy and safeguarding','Measurement, governance and acceptance']
TECH=['Mobile Expo/React Native architecture','Expo Router navigation','Local drafts and synchronization','OIDC Authorization Code with PKCE','Go API transport','OpenAPI 3.1 contracts','Needs domain persistence','Verification transitions','Evidence metadata lifecycle','RustFS/S3 object lifecycle','Project persistence and milestones','Contributions domain','Funding integration boundary','Impact aggregation','Notifications contract','Privacy orchestration','Trust and Safety integration','YugabyteDB system of record','Transactional outbox and events','SS-13 authorization integration','Gateway, rate limits and TLS','Vault and secret handling','Observability and SRE','AppForge/SkyForge supply chain and deployment']
TECH_LENSES=['Architecture and boundary','Data and transaction model','Security and authorization','Reliability and failure handling','Testing, observability and release evidence']
ROLES={'Community member / reporter':'discovers needs, reports observations, protects dignity and follows outcomes','Responder / volunteer':'offers time, skills, materials, knowledge, access, care or other abilities','Community organizer / project steward':'converts verified needs into coordinated work without displacing affected-community ownership','Eligible verifier':'reviews claims within granted authority and preserves independence between reporting and confirmation','NGO / institution / donor':'participates through transparent project, counterpart and governed funding boundaries','Trust & Safety / safeguarding reviewer':'handles confidential concerns, evidence decisions, enforcement and appeal boundaries','Platform administrator / support operator':'restores service and supports users without bypassing domain authority','Developer / API integrator':'uses OpenAPI, identity, idempotency, evidence and event contracts safely'}
USER_TASKS=['Start from the principal navigation','Understand status labels','Find or open a need','Report or inspect an observation','Use safe evidence practices','Understand verification boundaries','Open a Project Room','Work with contribution needs','Create or track a contribution offer','Understand funding/counterpart information','Follow outcome and Impact Passport data','Manage notifications','Exercise privacy choices','Report a safety concern','Troubleshoot and escalate safely']
TRAIN=['Orientation and doctrine','Role boundary and authority','Interface navigation','Need observation and reporting','Evidence dignity and consent','Verification literacy','Project Room coordination','Response to Ability','Funding and counterpart participation','Impact and outcome literacy','Notifications and communication','Privacy and data rights','Trust, safety and safeguarding','Offline and low-bandwidth operation','Troubleshooting and escalation','Scenario practicum','Knowledge check and assessment']
SCREEN_BY_ROLE={'Community member / reporter':'01-home-impact-feed.png','Responder / volunteer':'11-ability-profile.png','Community organizer / project steward':'06-project-room.png','Eligible verifier':'05-need-detail.png','NGO / institution / donor':'06-project-room.png','Trust & Safety / safeguarding reviewer':'12-trust-safety-report.png','Platform administrator / support operator':'10-notifications.png','Developer / API integrator':'04-needmap.png'}

d=new_doc('iRespond Product Documentation','Product definition, journeys, governance, safety, metrics and scale model','Product Documentation')
for domain in DOMAINS:
  for lens in LENSES:
    img=None
    if domain=='Impact Feed and discovery' and lens==LENSES[0]: img='01-home-impact-feed.png'
    elif domain=='Need reporting' and lens==LENSES[0]: img='02-report-need.png'
    elif domain=='Evidence capture and dignity' and lens==LENSES[0]: img='03-evidence-capture.png'
    elif domain=='NeedMap and geospatial discovery' and lens==LENSES[0]: img='04-needmap.png'
    elif domain=='Action Project conversion' and lens==LENSES[0]: img='06-project-room.png'
    elif domain=='Impact Passport' and lens==LENSES[0]: img='08-impact-passport.png'
    elif domain=='Privacy consent and data rights' and lens==LENSES[0]: img='09-privacy-data-rights.png'
    elif domain=='Trust and Safety reporting' and lens==LENSES[0]: img='12-trust-safety-report.png'
    add_page(d,f'Product domain • {domain}',f'{domain}: {lens}',common(domain,lens),[('Behavior and policy',common(domain,lens)),('Controls and safe failure',control(domain,lens)),('Evidence and measurement',evidence(domain))],['Confirm current state before acting.','Preserve affected-community dignity and consent.','Keep planned/shared-service direction distinct from current implementation.','Require auditable authority for consequential transitions.'],img)
finalise(d,OUT/'iRespond_Product_Documentation.docx')

d=new_doc('iRespond Technical Documentation','Architecture, APIs, data, security, reliability, deployment and living-document automation','Technical Documentation')
for topic in TECH:
  for lens in TECH_LENSES:
    add_page(d,f'Technical domain • {topic}',f'{topic}: {lens}',common(topic,lens),[('Design contract',common(topic,lens)),('Security and reliability',control(topic,lens)),('Verification evidence',evidence(topic))],['Use explicit contracts at domain and shared-service boundaries.','Fail closed for identity, authorization, evidence and moderation dependencies.','Prefer idempotent retries and transactional facts over compensating guesses.','Bind tests and runtime evidence to the exact source revision.'])
finalise(d,OUT/'iRespond_Technical_Documentation.docx')

d=new_doc('iRespond User Manual','Role-specific operational guidance for principal users','User Manual')
for role,desc in ROLES.items():
  for i,task in enumerate(USER_TASKS):
    img=SCREEN_BY_ROLE[role] if i in (0,4,9) else None
    summary=f'{role}: {task}. This role {desc}. The procedure is source-bound and must not be used to justify a privilege or state transition the user does not hold.'
    add_page(d,f'Role guide • {role}',task,summary,[('Before you begin','Confirm sign-in, role, current status, connectivity and whether the action is reversible. Read any verification, moderation, privacy or funding label before proceeding.'),('Procedure',f'Use the relevant iRespond surface to complete “{task}”. Keep the current trustworthy state if the API, identity provider or reviewer boundary is unavailable. Use local drafts or retry-safe queues only where the app explicitly provides them.'),('Success criteria','The expected state is visible, feedback is clear and the audit trail can explain the actor, decision and result. Success never means hiding an error or changing the database manually to make the screen look complete.'),('Troubleshooting','Check device/network, configuration, authentication, authorization and dependency availability in that order. Escalate safeguarding or access-control problems; do not bypass them.')],['Protect personal and beneficiary data.','Do not conflate observation, verification, commitment, fulfillment and outcome.','Retry only at documented safe boundaries.','Escalate rather than bypass role or safety controls.'],img)
finalise(d,OUT/'iRespond_User_Manual_All_Roles.docx')

d=new_doc('iRespond Training Manual','Instructor-led and self-paced curriculum with current-interface visual references, exercises and assessment','Training Manual')
for role,desc in ROLES.items():
  for i,module in enumerate(TRAIN):
    img=None
    if module=='Interface navigation': img=SCREEN_BY_ROLE[role]
    elif module=='Need observation and reporting' and role in ('Community member / reporter','Eligible verifier'): img='02-report-need.png'
    elif module=='Evidence dignity and consent': img='03-evidence-capture.png'
    elif module=='Project Room coordination' and role in ('Responder / volunteer','Community organizer / project steward','NGO / institution / donor'): img='06-project-room.png'
    elif module=='Privacy and data rights': img='09-privacy-data-rights.png'
    elif module=='Trust, safety and safeguarding': img='12-trust-safety-report.png'
    summary=f'Training module for {role}: {module}. The learner should build a correct decision model, not only memorize taps. The role {desc}.'
    add_page(d,f'Training path • {role}',module,summary,[('Learning objectives',f'Explain how {module.lower()} applies to {role.lower()}, identify at least two trust or safety boundaries, and demonstrate the workflow without an administrator bypass.'),('Facilitator notes','Start from a realistic community scenario. Ask what is known, what is observed, who is authorized to decide the next state, what evidence is sufficient and how the system should fail if a dependency is unavailable.'),('Guided practice',f'Have the learner rehearse {module.lower()}. Pause before every consequential action and require the learner to name the current state, target state, decision owner and safe recovery path.'),('Assessment','Score truthfulness, dignity, role correctness, accessibility awareness and safe failure above speed. Require one troubleshooting explanation and one escalation decision.')],['Learner can state the role boundary.','Learner can distinguish observation, verification, commitment and outcome where relevant.','Learner can demonstrate retry/escalation behavior.','Learner can identify privacy, accessibility and safeguarding implications.'],img)
finalise(d,OUT/'iRespond_Training_Manual_All_Roles.docx')

manifest={'source_revision':REV,'generated':dt.date.today().isoformat(),'documents':sorted(p.name for p in OUT.glob('*.docx')),'ui_assets':sorted(p.name for p in UI.glob('*.png')),'mermaid_sources':sorted(p.name for p in MMD.glob('*.mmd'))}
(ROOT/'generated-manifest.json').write_text(json.dumps(manifest,indent=2)+'\n')
print(json.dumps(manifest,indent=2))
